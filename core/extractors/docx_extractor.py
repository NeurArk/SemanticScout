from __future__ import annotations

import logging
from typing import Any, Dict

from docx import Document as DocxDocument

from core.exceptions import DocumentProcessingError

from .base_extractor import BaseExtractor

logger = logging.getLogger(__name__)


class DOCXExtractor(BaseExtractor):
    """Extract text from Word documents."""

    def can_extract(self, file_path: str) -> bool:
        return file_path.lower().endswith(".docx")

    def extract(self, file_path: str) -> Dict[str, Any]:
        """Extract text and metadata from DOCX."""
        self.validate_file(file_path)

        try:
            doc = DocxDocument(file_path)

            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

            table_texts: list[str] = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    if any(row_text):
                        table_text.append(" | ".join(row_text))
                if table_text:
                    table_texts.append("\n".join(table_text))

            full_text = "\n\n".join(paragraphs)
            if table_texts:
                full_text += "\n\n[Tables]\n" + "\n\n".join(table_texts)

            metadata = {
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
                "author": doc.core_properties.author or "",
                "title": doc.core_properties.title or "",
                "created": str(doc.core_properties.created) if doc.core_properties.created else "",
            }

            return {"content": full_text, "metadata": metadata}
        except Exception as exc:  # pragma: no cover - external lib
            logger.error("DOCX extraction failed: %s", exc)
            raise DocumentProcessingError(f"Failed to extract DOCX: {exc}") from exc
